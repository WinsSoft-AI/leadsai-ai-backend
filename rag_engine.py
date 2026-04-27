"""
Winssoft BMA — RAG Engine v2
Multi-strategy retrieval: ChromaDB (vector) + BM25 (keyword) + RRF fusion

Vector DB: ChromaDB (open-source, no API key needed)
  - Runs in-memory (dev), persistent on disk (default), or as HTTP server (production cluster)
  - pip install chromadb

Fallback: if chromadb not installed, falls back to in-memory cosine similarity
"""
import os
import re
import json
import math
import uuid
import asyncio
import hashlib
import logging
from typing import List, Dict, Any, Optional, Tuple
from collections import defaultdict

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────
# ChromaDB import with graceful fallback
# ──────────────────────────────────────────────────────────────
try:
    import chromadb
    from chromadb.config import Settings as ChromaSettings
    CHROMA_AVAILABLE = True
    logger.info("✅ ChromaDB available")
except ImportError:
    CHROMA_AVAILABLE = False
    logger.warning("⚠️  ChromaDB not installed — using in-memory fallback. Run: pip install chromadb")

# PDF parsing — try pypdf first, then PyPDF2, then basic fallback
try:
    from pypdf import PdfReader as _PdfReader
    PDF_LIB = "pypdf"
except ImportError:
    try:
        from PyPDF2 import PdfReader as _PdfReader
        PDF_LIB = "pypdf2"
    except ImportError:
        _PdfReader = None
        PDF_LIB = "basic"

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# BeautifulSoup for HTML
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


# ══════════════════════════════════════════════════════════════
# TEXT CHUNKER
# ══════════════════════════════════════════════════════════════
class TextChunker:
    def __init__(self, chunk_size: int = 400, overlap: int = 50):
        self.chunk_size  = chunk_size
        self.overlap     = overlap

    def chunk(self, text: str, source: str = "document") -> List[Dict]:
        """Split text into overlapping chunks with metadata."""
        text  = re.sub(r'\s+', ' ', text).strip()
        words = text.split()
        chunks, i, idx = [], 0, 0

        while i < len(words):
            chunk_words = words[i : i + self.chunk_size]
            chunk_text  = " ".join(chunk_words)

            if len(chunk_text.strip()) > 50:
                uid = hashlib.md5(f"{source}:{i}:{chunk_text[:80]}".encode()).hexdigest()[:12]
                chunks.append({
                    "id":          uid,
                    "text":        chunk_text,
                    "source":      source,
                    "chunk_index": idx,
                    "word_count":  len(chunk_words),
                })
                idx += 1

            i += self.chunk_size - self.overlap

        return chunks


# ══════════════════════════════════════════════════════════════
# CHROMA VECTOR STORE
# ══════════════════════════════════════════════════════════════
class ChromaVectorStore:
    """
    ChromaDB-backed vector store.

    Modes (set via config):
      memory     → in-process ephemeral (dev/testing)
      persistent → saves to disk at chroma_persist_dir (default production)
      http       → connects to a remote ChromaDB server
    """

    def __init__(self, mode: str = "persistent", persist_dir: str = "./chroma_data",
                 host: str = "localhost", port: int = 8001):
        self.mode = mode
        self._client = None
        self._collections: Dict[str, Any] = {}

        if not CHROMA_AVAILABLE:
            logger.warning("ChromaDB unavailable — using in-memory fallback")
            self._fallback: Dict[str, List[Dict]] = {}
            return

        try:
            if mode == "memory":
                self._client = chromadb.EphemeralClient()
            elif mode == "http":
                self._client = chromadb.HttpClient(host=host, port=port)
            else:  # persistent (default)
                os.makedirs(persist_dir, exist_ok=True)
                self._client = chromadb.PersistentClient(path=persist_dir)

            logger.info(f"✅ ChromaDB client ready (mode={mode})")
        except Exception as e:
            logger.error(f"ChromaDB init failed: {e} — falling back to in-memory")
            self._client = None
            self._fallback = {}

    def _get_collection(self, tenant_id: str):
        """Get or create a ChromaDB collection for a tenant."""
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', tenant_id)[:63]
        if safe_name not in self._collections:
            self._collections[safe_name] = self._client.get_or_create_collection(
                name=safe_name,
                metadata={"hnsw:space": "cosine"},  # cosine similarity
            )
        return self._collections[safe_name]

    # ── Upsert ────────────────────────────────────────────────
    def upsert(self, tenant_id: str, chunks: List[Dict]):
        if not chunks:
            return

        if self._client is None:
            # Fallback: plain list
            if tenant_id not in self._fallback:
                self._fallback[tenant_id] = []
            existing_ids = {c["id"] for c in self._fallback[tenant_id]}
            for c in chunks:
                if c["id"] not in existing_ids:
                    self._fallback[tenant_id].append(c)
            return

        col = self._get_collection(tenant_id)

        ids        = [c["id"]  for c in chunks]
        documents  = [c["text"] for c in chunks]
        embeddings = [c.get("embedding") for c in chunks]
        metadatas  = [{
            "source":      c.get("source",      "unknown"),
            "chunk_index": c.get("chunk_index", 0),
            "word_count":  c.get("word_count",  0),
        } for c in chunks]

        if all(e is not None for e in embeddings):
            col.upsert(ids=ids, documents=documents, embeddings=embeddings, metadatas=metadatas)
        else:
            # No embeddings yet — store without; similarity search won't work until embedded
            col.upsert(ids=ids, documents=documents, metadatas=metadatas)

    # ── Search ────────────────────────────────────────────────
    def search(self, tenant_id: str, query_vector: List[float], top_k: int = 10) -> List[Dict]:
        if self._client is None:
            return self._fallback_search(tenant_id, query_vector, top_k)

        try:
            col = self._get_collection(tenant_id)
            if col.count() == 0:
                return []

            results = col.query(
                query_embeddings=[query_vector],
                n_results=min(top_k, col.count()),
                include=["documents", "metadatas", "distances"],
            )

            out = []
            for i, doc_id in enumerate(results["ids"][0]):
                out.append({
                    "id":          doc_id,
                    "text":        results["documents"][0][i],
                    "source":      results["metadatas"][0][i].get("source", ""),
                    "chunk_index": results["metadatas"][0][i].get("chunk_index", 0),
                    # ChromaDB returns distance (lower=better); convert to score
                    "score":       1.0 - results["distances"][0][i],
                })
            return out

        except Exception as e:
            logger.error(f"ChromaDB search error: {e}")
            return []

    # ── Get all chunks ────────────────────────────────────────
    def get_all_chunks(self, tenant_id: str) -> List[Dict]:
        if self._client is None:
            return self._fallback.get(tenant_id, [])
        try:
            col = self._get_collection(tenant_id)
            if col.count() == 0:
                return []
            res = col.get(include=["documents", "metadatas"])
            return [
                {"id": res["ids"][i], "text": res["documents"][i], **res["metadatas"][i]}
                for i in range(len(res["ids"]))
            ]
        except Exception:
            return []

    # ── Count ─────────────────────────────────────────────────
    def count(self, tenant_id: str) -> int:
        if self._client is None:
            return len(self._fallback.get(tenant_id, []))
        try:
            return self._get_collection(tenant_id).count()
        except Exception:
            return 0

    # ── Delete by IDs ─────────────────────────────────────────
    def delete_by_ids(self, tenant_id: str, ids: List[str]):
        """Remove specific chunks by their IDs from a tenant's collection."""
        if not ids:
            return
        if self._client is None:
            if tenant_id in self._fallback:
                self._fallback[tenant_id] = [
                    c for c in self._fallback[tenant_id] if c["id"] not in ids
                ]
            return
        try:
            col = self._get_collection(tenant_id)
            col.delete(ids=ids)
        except Exception as e:
            logger.error(f"ChromaDB delete_by_ids error: {e}")

    # ── Delete tenant ─────────────────────────────────────────
    def delete_tenant(self, tenant_id: str):
        if self._client is None:
            self._fallback.pop(tenant_id, None)
            return
        try:
            safe = re.sub(r'[^a-zA-Z0-9_-]', '_', tenant_id)[:63]
            self._client.delete_collection(safe)
            self._collections.pop(safe, None)
        except Exception:
            pass

    # ── In-memory cosine fallback ─────────────────────────────
    def _fallback_search(self, tenant_id: str, query_vec: List[float], top_k: int) -> List[Dict]:
        chunks = self._fallback.get(tenant_id, [])
        results = []
        for c in chunks:
            emb = c.get("embedding", [])
            if emb:
                score = _cosine_similarity(query_vec, emb)
                results.append({**c, "score": score})
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]


def _cosine_similarity(a: List[float], b: List[float]) -> float:
    if len(a) != len(b) or not a:
        return 0.0
    dot    = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x ** 2 for x in a))
    norm_b = math.sqrt(sum(x ** 2 for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


# ══════════════════════════════════════════════════════════════
# BM25 SPARSE RETRIEVER
# ══════════════════════════════════════════════════════════════
class BM25Retriever:
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1     = k1
        self.b      = b
        self._index: Dict[str, Dict] = {}

    def index(self, tenant_id: str, chunks: List[Dict]):
        if not chunks:
            return
        tokenized = [self._tokenize(c["text"]) for c in chunks]
        df: Dict[str, int] = defaultdict(int)
        for tokens in tokenized:
            for token in set(tokens):
                df[token] += 1
        avg_len = sum(len(t) for t in tokenized) / len(tokenized)
        self._index[tenant_id] = {
            "chunks":    chunks,
            "tokenized": tokenized,
            "df":        dict(df),
            "avg_len":   avg_len,
            "n":         len(chunks),
        }

    def search(self, tenant_id: str, query: str, top_k: int = 10) -> List[Dict]:
        idx = self._index.get(tenant_id)
        if not idx:
            return []
        query_tokens = self._tokenize(query)
        scores = []
        for chunk, tokens in zip(idx["chunks"], idx["tokenized"]):
            score  = 0.0
            tf_map = defaultdict(int)
            for t in tokens:
                tf_map[t] += 1
            for token in query_tokens:
                if token not in idx["df"]:
                    continue
                tf   = tf_map.get(token, 0)
                df   = idx["df"][token]
                n    = idx["n"]
                idf  = math.log((n - df + 0.5) / (df + 0.5) + 1)
                tf_n = (tf * (self.k1 + 1)) / (
                    tf + self.k1 * (1 - self.b + self.b * len(tokens) / max(idx["avg_len"], 1))
                )
                score += idf * tf_n
            if score > 0:
                scores.append({**chunk, "bm25_score": score})
        scores.sort(key=lambda x: x["bm25_score"], reverse=True)
        return scores[:top_k]

    def _tokenize(self, text: str) -> List[str]:
        text  = re.sub(r'[^a-z0-9\s]', ' ', text.lower())
        stops = {
            'the','a','an','is','are','was','were','be','been','being',
            'have','has','had','do','does','did','will','would','could',
            'should','may','might','to','of','in','for','on','with','at',
            'by','from','as','into','and','or','but','if','it','its',
            'this','that','these','those','not','no','so','we','he','she',
            'they','their','our','your','i','me','my','you',
        }
        return [t for t in text.split() if t not in stops and len(t) > 1]


# ══════════════════════════════════════════════════════════════
# DOCUMENT PARSER
# ══════════════════════════════════════════════════════════════
class DocumentParser:

    def parse(self, content: bytes, filename: str, content_type: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext in ("txt",) or content_type == "text/plain":
            return content.decode("utf-8", errors="ignore")

        elif ext in ("html", "htm") or content_type == "text/html":
            return self._parse_html(content.decode("utf-8", errors="ignore"))

        elif ext == "md":
            return self._parse_markdown(content.decode("utf-8", errors="ignore"))

        elif ext == "json":
            try:
                return self._flatten_json(json.loads(content))
            except Exception:
                return content.decode("utf-8", errors="ignore")

        elif ext == "csv":
            return self._parse_csv(content.decode("utf-8", errors="ignore"))

        elif ext == "pdf":
            return self._parse_pdf(content)

        elif ext in ("docx", "doc"):
            return self._parse_docx(content)

        else:
            return content.decode("utf-8", errors="ignore")

    # ── PDF ───────────────────────────────────────────────────
    def _parse_pdf(self, content: bytes) -> str:
        if _PdfReader is not None:
            try:
                import io
                reader = _PdfReader(io.BytesIO(content))
                parts  = []
                for page in reader.pages:
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        pass
                text = "\n".join(parts)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PDF lib parse failed: {e}")

        # Basic BT/ET regex fallback
        try:
            raw = content.decode("latin-1", errors="ignore")
            blocks = re.findall(r'BT(.*?)ET', raw, re.DOTALL)
            parts  = []
            for blk in blocks:
                parts.extend(re.findall(r'\(([^)]*)\)\s*Tj', blk))
                for arr in re.findall(r'\[([^\]]*)\]\s*TJ', blk):
                    parts.extend(re.findall(r'\(([^)]*)\)', arr))
            return " ".join(parts)
        except Exception:
            return ""

    # ── DOCX ──────────────────────────────────────────────────
    def _parse_docx(self, content: bytes) -> str:
        if DOCX_AVAILABLE:
            try:
                import io
                doc   = DocxDocument(io.BytesIO(content))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also grab table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                parts.append(cell.text.strip())
                return "\n".join(parts)
            except Exception as e:
                logger.warning(f"DOCX parse failed: {e}")
        return content.decode("utf-8", errors="ignore")

    # ── HTML ──────────────────────────────────────────────────
    def _parse_html(self, html: str) -> str:
        if BS4_AVAILABLE:
            try:
                soup = BeautifulSoup(html, "lxml")
                for tag in soup(["script", "style", "nav", "footer", "header"]):
                    tag.decompose()
                return soup.get_text(separator=" ", strip=True)
            except Exception:
                pass
        # Regex fallback
        html = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<[^>]+>', ' ', html)
        html = re.sub(r'&(?:nbsp|amp|lt|gt|quot);', lambda m: {
            '&nbsp;':' ','&amp;':'&','&lt;':'<','&gt;':'>','&quot;':'"'
        }.get(m.group(),' '), html)
        return re.sub(r'\s+', ' ', html).strip()

    # ── Markdown ──────────────────────────────────────────────
    def _parse_markdown(self, text: str) -> str:
        text = re.sub(r'#{1,6}\s', '', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',   r'\1', text)
        text = re.sub(r'`(.+?)`',     r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'!\[.*?\]\(.+?\)', '', text)
        return text

    # ── CSV ───────────────────────────────────────────────────
    def _parse_csv(self, csv_text: str) -> str:
        lines   = csv_text.strip().splitlines()
        if not lines:
            return ""
        headers = [h.strip().strip('"') for h in lines[0].split(',')]
        rows    = []
        for line in lines[1:]:
            vals    = [v.strip().strip('"') for v in line.split(',')]
            row_txt = " | ".join(f"{h}: {v}" for h, v in zip(headers, vals) if v)
            if row_txt:
                rows.append(row_txt)
        return "\n".join(rows)

    # ── JSON ──────────────────────────────────────────────────
    def _flatten_json(self, data: Any, prefix: str = "") -> str:
        parts = []
        if isinstance(data, dict):
            for k, v in data.items():
                parts.append(self._flatten_json(v, f"{prefix}.{k}" if prefix else k))
        elif isinstance(data, list):
            for i, item in enumerate(data):
                parts.append(self._flatten_json(item, f"{prefix}[{i}]"))
        else:
            parts.append(f"{prefix}: {data}" if prefix else str(data))
        return " ".join(filter(None, parts))


# ══════════════════════════════════════════════════════════════
# MAIN RAG ENGINE
# ══════════════════════════════════════════════════════════════
class RAGEngine:
    def __init__(self, chroma_mode: str = "persistent",
                 chroma_persist_dir: str = "./chroma_data",
                 chroma_host: str = "localhost", chroma_port: int = 8001):

        self.chunker      = TextChunker(chunk_size=400, overlap=50)
        self.vector_store = ChromaVectorStore(
            mode=chroma_mode,
            persist_dir=chroma_persist_dir,
            host=chroma_host,
            port=chroma_port,
        )
        self.bm25    = BM25Retriever()
        self.parser  = DocumentParser()
        self._jobs:   Dict[str, Dict] = {}
        self._gemini: Optional[Any]   = None

    def set_gemini(self, gemini_client):
        self._gemini = gemini_client

    # ── Ingest ────────────────────────────────────────────────
    async def ingest(self, content: bytes, filename: str,
                     content_type: str, tenant_id: str, pool) -> str:
        job_id = uuid.uuid4().hex[:8]
        self._jobs[job_id] = {
            "job_id":         job_id,
            "status":         "processing",
            "filename":       filename,
            "chunks_indexed": 0,
            "error":          None,
        }
        async with pool.acquire() as conn:
            await conn.execute(
                "INSERT INTO ingest_jobs (id, tenant_id, filename, file_size_bytes, status) VALUES ($1, $2, $3, $4, $5)",
                job_id, tenant_id, filename, len(content), "processing"
            )
        asyncio.create_task(
            self._ingest_task(job_id, content, filename, content_type, tenant_id, pool)
        )
        return job_id

    async def _ingest_task(self, job_id, content, filename, content_type, tenant_id, pool):
        try:
            text = self.parser.parse(content, filename, content_type)
            if not text.strip():
                self._jobs[job_id].update(status="failed", error="No extractable text")
                async with pool.acquire() as conn:
                    await conn.execute("UPDATE ingest_jobs SET status='failed', error_message=$1 WHERE id=$2", "No extractable text", job_id)
                return

            chunks = self.chunker.chunk(text, source=filename)
            for c in chunks:
                c["id"] = f"{job_id}_{c['id']}"
            if not chunks:
                self._jobs[job_id].update(status="failed", error="No chunks produced")
                async with pool.acquire() as conn:
                    await conn.execute("UPDATE ingest_jobs SET status='failed', error_message=$1 WHERE id=$2", "No chunks produced", job_id)
                return

            # Embed in batches of 10
            if self._gemini:
                texts = [c["text"] for c in chunks]
                all_emb = []
                for i in range(0, len(texts), 10):
                    batch_emb = await self._gemini.get_embeddings(texts[i:i+10])
                    all_emb.extend(batch_emb)
                for chunk, emb in zip(chunks, all_emb):
                    chunk["embedding"] = emb

            self.vector_store.upsert(tenant_id, chunks)

            # Rebuild BM25 index
            all_chunks = self.vector_store.get_all_chunks(tenant_id)
            self.bm25.index(tenant_id, all_chunks)

            self._jobs[job_id].update(status="completed", chunks_indexed=len(chunks))
            async with pool.acquire() as conn:
                await conn.execute("UPDATE ingest_jobs SET status='completed', chunks_indexed=$1, completed_at=NOW() WHERE id=$2", len(chunks), job_id)
            logger.info(f"✅ Ingested {filename}: {len(chunks)} chunks → tenant {tenant_id}")

        except Exception as e:
            logger.error(f"Ingest task failed: {e}", exc_info=True)
            self._jobs[job_id].update(status="failed", error=str(e))
            async with pool.acquire() as conn:
                await conn.execute("UPDATE ingest_jobs SET status='failed', error_message=$1 WHERE id=$2", str(e), job_id)

    # ── Retrieve ──────────────────────────────────────────────
    async def retrieve(self, query: str, tenant_id: str, top_k: int = 5) -> List[Dict]:
        """Hybrid: ChromaDB vector search + BM25 keyword + exact text match, fused with RRF."""
        query_vec = []
        if self._gemini:
            try:
                embs      = await self._gemini.get_embeddings([query])
                query_vec = embs[0] if embs else []
            except Exception as e:
                logger.warning(f"Embedding error: {e}")

        # 1. Vector similarity search (broad)
        vector_results = []
        if query_vec:
            vector_results = self.vector_store.search(tenant_id, query_vec, top_k=top_k * 3)

        # 2. BM25 keyword search (broad)
        bm25_results = self.bm25.search(tenant_id, query, top_k=top_k * 3)

        # 3. Direct text search in ChromaDB for exact matches the embeddings might miss
        exact_results = []
        try:
            if self.vector_store._client is not None:
                col = self.vector_store._get_collection(tenant_id)
                if col.count() > 0:
                    # Search for documents containing key query words
                    query_words = [w for w in query.lower().split() if len(w) > 2]
                    for word in query_words[:3]:  # limit to 3 keywords
                        try:
                            text_hits = col.get(
                                where_document={"$contains": word},
                                include=["documents", "metadatas"],
                                limit=top_k,
                            )
                            if text_hits and text_hits["ids"]:
                                for i, doc_id in enumerate(text_hits["ids"]):
                                    exact_results.append({
                                        "id": doc_id,
                                        "text": text_hits["documents"][i],
                                        "source": text_hits["metadatas"][i].get("source", ""),
                                        "chunk_index": text_hits["metadatas"][i].get("chunk_index", 0),
                                        "score": 0.85,  # give a solid base score for exact text hits
                                    })
                        except Exception:
                            pass
        except Exception as e:
            logger.warning(f"Exact text search error: {e}")

        return self._rrf_fusion(vector_results, bm25_results, exact_results, top_k=top_k)

    def _rrf_fusion(self, list1: List[Dict], list2: List[Dict],
                    list3: List[Dict] = None,
                    k: int = 60, top_k: int = 5) -> List[Dict]:
        scores: Dict[str, float] = {}
        docs:   Dict[str, Dict]  = {}

        for rank, doc in enumerate(list1):
            did = doc["id"]
            scores[did] = scores.get(did, 0) + 1 / (k + rank + 1)
            docs[did]   = doc

        for rank, doc in enumerate(list2):
            did = doc["id"]
            scores[did] = scores.get(did, 0) + 1 / (k + rank + 1)
            docs[did]   = doc

        if list3:
            for rank, doc in enumerate(list3):
                did = doc["id"]
                scores[did] = scores.get(did, 0) + 1 / (k + rank + 1)
                if did not in docs:
                    docs[did] = doc

        result = []
        for did in sorted(scores, key=lambda x: scores[x], reverse=True)[:top_k]:
            d = {k_: v for k_, v in docs[did].items() if k_ != "embedding"}
            d["score"] = scores[did]
            result.append(d)

        return result

    # ── Custom Q/A Ingestion ───────────────────────────────────
    async def ingest_qa(self, qa_id: str, question: str, answer: str, tenant_id: str):
        """
        Ingest a single custom Q/A pair as one dedicated chunk.
        Uses a 'qa_' prefix so it can be identified and deleted later.
        """
        chunk_id = f"qa_{qa_id}"
        text = f"Question: {question}\nAnswer: {answer}"
        chunk = {
            "id":          chunk_id,
            "text":        text,
            "source":      "Custom Q/A",
            "chunk_index": 0,
            "word_count":  len(text.split()),
        }

        # Embed
        if self._gemini:
            try:
                embs = await self._gemini.get_embeddings([text])
                if embs:
                    chunk["embedding"] = embs[0]
            except Exception as e:
                logger.warning(f"QA embedding error: {e}")

        self.vector_store.upsert(tenant_id, [chunk])

        # Rebuild BM25
        all_chunks = self.vector_store.get_all_chunks(tenant_id)
        self.bm25.index(tenant_id, all_chunks)
        logger.info(f"✅ Indexed Q/A '{qa_id}' for tenant {tenant_id}")

    async def delete_qa_chunk(self, qa_id: str, tenant_id: str):
        """Remove a custom Q/A chunk from the vector store."""
        chunk_id = f"qa_{qa_id}"
        self.vector_store.delete_by_ids(tenant_id, [chunk_id])

        # Rebuild BM25
        all_chunks = self.vector_store.get_all_chunks(tenant_id)
        self.bm25.index(tenant_id, all_chunks)
        logger.info(f"🗑️ Deleted Q/A chunk '{qa_id}' for tenant {tenant_id}")

    async def delete_doc_chunks(self, doc_id: str, tenant_id: str):
        """Remove all chunks for a specific document (job_id) from the vector store."""
        all_chunks = self.vector_store.get_all_chunks(tenant_id)
        ids_to_delete = [c["id"] for c in all_chunks if c["id"].startswith(f"{doc_id}_")]
        if ids_to_delete:
            self.vector_store.delete_by_ids(tenant_id, ids_to_delete)
            # Rebuild BM25
            remaining = self.vector_store.get_all_chunks(tenant_id)
            self.bm25.index(tenant_id, remaining)
            logger.info(f"🗑️ Deleted {len(ids_to_delete)} chunks for doc '{doc_id}' in tenant {tenant_id}")

    # ── Rebuild Structured KB ─────────────────────────────────
    async def rebuild_structured(self, tenant_id: str, chunks: List[Dict]):
        """
        Flush all structured KB chunks (kb_company_*, kb_product_*, qa_*) for a tenant,
        then re-embed and upsert the provided chunks. File-upload chunks are preserved.
        """
        # 1. Get all existing chunks and identify structured-KB ones
        all_existing = self.vector_store.get_all_chunks(tenant_id)
        kb_prefixes = ("kb_company_", "kb_product_", "qa_")
        ids_to_delete = [c["id"] for c in all_existing if c["id"].startswith(kb_prefixes)]

        if ids_to_delete:
            self.vector_store.delete_by_ids(tenant_id, ids_to_delete)
            logger.info(f"🗑️ Flushed {len(ids_to_delete)} structured KB chunks for tenant {tenant_id}")

        if not chunks:
            # Still rebuild BM25 from remaining file-upload chunks
            remaining = self.vector_store.get_all_chunks(tenant_id)
            self.bm25.index(tenant_id, remaining)
            return {"chunks_indexed": 0, "chunks_deleted": len(ids_to_delete)}

        # 2. Embed in batches of 10
        if self._gemini:
            texts = [c["text"] for c in chunks]
            all_emb = []
            for i in range(0, len(texts), 10):
                batch_emb = await self._gemini.get_embeddings(texts[i:i+10])
                all_emb.extend(batch_emb)
            for chunk, emb in zip(chunks, all_emb):
                chunk["embedding"] = emb

        # 3. Upsert new structured chunks
        self.vector_store.upsert(tenant_id, chunks)

        # 4. Rebuild BM25 from all chunks (file + structured)
        all_chunks = self.vector_store.get_all_chunks(tenant_id)
        self.bm25.index(tenant_id, all_chunks)

        logger.info(f"✅ Rebuilt structured KB for tenant {tenant_id}: {len(chunks)} chunks indexed")
        return {"chunks_indexed": len(chunks), "chunks_deleted": len(ids_to_delete)}

    # ── Helpers ───────────────────────────────────────────────
    async def get_job_status(self, job_id: str) -> Dict:
        return self._jobs.get(job_id, {"status": "not_found"})

    def chunk_count(self, tenant_id: str) -> int:
        return self.vector_store.count(tenant_id)

